#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_doc.py — tangoya 바이브코딩 수업용 문서 생성기
python3 build/make_doc.py → tangoya_vibecoding.docx 생성
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT  = os.path.join(ROOT, 'tangoya_vibecoding.docx')

# ──────────────────────────────────────────────
# 헬퍼 함수
# ──────────────────────────────────────────────
def set_font(run, name_ko='맑은 고딕', name_en='Calibri', size=None, bold=False,
             color=None, italic=False):
    run.font.name = name_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), name_ko)
    if size:  run.font.size = Pt(size)
    if bold:  run.font.bold = True
    if italic: run.font.italic = True
    if color: run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1, color=(0,0,0)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Calibri'
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
        run.font.color.rgb = RGBColor(*color)
    return p


def add_para(doc, text, size=11, bold=False, color=(50,50,50),
             indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent: p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_code_block(doc, code_text):
    """회색 배경 코드 블록"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        # 회색 단락 배경
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(40, 40, 40)
    doc.add_paragraph()


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 헤더
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
        # 헤더 배경색
        tc_pr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2563EB')
        tc_pr.append(shd)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255,255,255)
    # 데이터
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
            if ri % 2 == 1:
                tc_pr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F0F4FF')
                tc_pr.append(shd)
    # 열 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_step_box(doc, step_num, title, desc):
    """스텝 박스 (파란 번호 + 제목 + 설명)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    # 번호
    r1 = p.add_run(f'  STEP {step_num}  ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(255,255,255)
    r1.font.name = 'Calibri'
    r1._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
    # 배경색 (단락 전체에)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1D4ED8')
    pPr.append(shd)
    # 제목
    r2 = p.add_run(f'  {title}')
    r2.font.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(255,255,255)
    r2.font.name = 'Calibri'
    r2._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
    # 설명 줄
    add_para(doc, desc, size=10, color=(80,80,80), indent=0.3, space_before=2, space_after=8)


def add_prompt_box(doc, prompt_text):
    """프롬프트 예시 박스 (노란 배경)"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FEF9C3')
    pPr.append(shd)
    p.paragraph_format.left_indent   = Cm(0.3)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run('💬 실제 프롬프트 예시')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(120, 80, 0)
    run.font.name = 'Calibri'
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')

    for line in prompt_text.strip().split('\n'):
        p2 = doc.add_paragraph()
        pPr2 = p2._element.get_or_add_pPr()
        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'), 'clear')
        shd2.set(qn('w:color'), 'auto')
        shd2.set(qn('w:fill'), 'FEF9C3')
        pPr2.append(shd2)
        p2.paragraph_format.left_indent  = Cm(0.6)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        run2 = p2.add_run(line if line else ' ')
        run2.font.size = Pt(9.5)
        run2.font.name = 'Courier New'
        run2._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
        run2.font.color.rgb = RGBColor(60, 40, 0)
    doc.add_paragraph()


def add_screenshot_placeholder(doc, caption):
    """스크린샷 자리 표시자"""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E5E7EB')
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ 스크린샷: {caption} ]')
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100,100,100)
    run.font.name = 'Calibri'
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
    doc.add_paragraph()


# ──────────────────────────────────────────────
# 문서 생성
# ──────────────────────────────────────────────
doc = Document()

# 기본 여백 설정
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ═══════════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run('바이브 코딩 실전 사례')
r.font.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(29, 78, 216)
r.font.name = 'Calibri'
r._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('tangoya (単語屋) 개발 과정')
r2.font.bold = True
r2.font.size = Pt(20)
r2.font.color.rgb = RGBColor(30,30,30)
r2.font.name = 'Calibri'
r2._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_desc.add_run('Claude Code와 함께하는 일본어 JLPT 레벨 판정 앱 제작 전 과정')
r3.font.size = Pt(13)
r3.font.color.rgb = RGBColor(80,80,80)
r3.font.name = 'Calibri'
r3._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()
doc.add_paragraph()

# 표지 정보 테이블
info_table = doc.add_table(rows=4, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ('앱 이름',    'tangoya (単語屋) — 일본어 JLPT 레벨 판정기'),
    ('개발자',     'Jaehyoring'),
    ('개발 도구',  'Claude Code (Anthropic) — 바이브 코딩 방식'),
    ('버전',       'v1.0  |  GitHub: github.com/Jaehyoring/tangoya'),
]
for ri, (k, v) in enumerate(info_data):
    c1, c2 = info_table.rows[ri].cells
    c1.text = k
    c2.text = v
    c1.paragraphs[0].runs[0].font.bold = True
    c1.paragraphs[0].runs[0].font.size = Pt(10)
    c1._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'DBEAFE')
    c1._element.tcPr.append(shd)
    c2.paragraphs[0].runs[0].font.size = Pt(10)
    for cell in (c1, c2):
        for run in cell.paragraphs[0].runs:
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '맑은 고딕')
    c1.width = Cm(3.5); c2.width = Cm(12)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. 바이브 코딩이란?
# ═══════════════════════════════════════════════
add_heading(doc, '1.  바이브 코딩(Vibe Coding)이란?', level=1, color=(29,78,216))

add_para(doc, '바이브 코딩(Vibe Coding)은 2025년 AI 연구자 Andrej Karpathy가 제안한 개발 방식으로, '
         '개발자가 코드를 직접 타이핑하는 대신 AI에게 자연어로 지시를 내려 소프트웨어를 만드는 방법입니다. '
         '"코드를 완전히 AI에게 맡기고, 개발자는 의도(vibe)만 전달한다"는 개념에서 이름이 붙었습니다.',
         size=11)

add_para(doc, '핵심 특징:', size=11, bold=True, space_before=6)

bullets = [
    '🗣️  자연어 프롬프트로 기능을 요청 — 코드 문법을 몰라도 됨',
    '🤖  AI(Claude Code 등)가 코드를 작성·수정·디버깅',
    '👁️  개발자는 결과를 확인하고 다음 방향을 지시',
    '🔄  대화 흐름으로 기능이 점진적으로 완성됨',
    '⚡  기획부터 배포까지 수 시간~수 일 안에 완성 가능',
]
for b in bullets:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(b)
    set_font(run, size=11, color=(40,40,40))

doc.add_paragraph()
add_para(doc, '전통적인 개발 방식과 비교:', size=11, bold=True, space_before=4)

add_table(doc,
    headers=['구분', '전통적 개발', '바이브 코딩'],
    rows=[
        ['코드 작성',   '개발자가 직접 타이핑',           'AI가 자동 생성'],
        ['필요 지식',   '언어 문법·라이브러리 숙지 필수',  '기능 요구사항 설명 능력'],
        ['오류 수정',   '디버깅 직접 수행',               'AI에게 오류 메시지 전달'],
        ['개발 속도',   '기능 1개에 수 시간~수 일',        '기능 1개에 수 분~수 시간'],
        ['생산성',      '개인 역량에 의존',               '아이디어가 곧 제품으로'],
    ],
    col_widths=[3, 6.5, 6.5]
)

add_heading(doc, '이 수업에서 배울 것', level=2, color=(55,65,81))
add_para(doc, 'tangoya 앱의 실제 개발 과정을 통해 바이브 코딩의 9단계 흐름을 체험합니다. '
         '프롬프트를 어떻게 작성하는지, AI와 어떻게 협업하는지, 완성된 앱을 어떻게 배포하는지를 '
         '단계별로 살펴봅니다.', size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 2. 프로젝트 소개
# ═══════════════════════════════════════════════
add_heading(doc, '2.  프로젝트 소개 — tangoya (単語屋)', level=1, color=(29,78,216))

add_para(doc, 'tangoya는 일본어 텍스트를 입력하면 각 단어의 JLPT(일본어능력시험) 레벨을 판정하고 '
         '한국어 뜻을 함께 보여주는 웹 앱입니다.', size=11)

add_table(doc,
    headers=['항목', '내용'],
    rows=[
        ['앱 이름',        'tangoya (単語屋) — 일본어로 "단어 가게"'],
        ['핵심 기능',      '일본어 입력 → 형태소 분석 → JLPT 레벨 표시 + 한국어 뜻'],
        ['JLPT 레벨',      'N1(최고급) ~ N5(초급) + 外(미등재) + 文法(문법요소)'],
        ['내장 단어 수',   '13,680개 항목 (N1~N5 7,680단어 × 히라가나+한자 이중 등록)'],
        ['한국어 뜻',      '7,518개 단어 자동 번역 수록'],
        ['배포 방식',      '단일 HTML 파일 + 로컬 서버 (설치 불필요)'],
        ['사용 기술',      'HTML5 · CSS3 · JavaScript · Kuromoji.js · Python'],
        ['부가 기능',      '다크/라이트 모드 · 관리자 편집 · 사전 커스텀 · JSON/CSV/TXT 내보내기'],
    ],
    col_widths=[4, 12]
)

add_screenshot_placeholder(doc, 'tangoya 앱 메인 화면 (다크모드)')
add_screenshot_placeholder(doc, '문장 분석 결과 — 토큰 카드 표시')

add_heading(doc, '완성된 앱의 주요 화면', level=2, color=(55,65,81))

screens = [
    ('단어 단독 입력 (Case A)',
     '단어 하나를 입력하면 대형 카드로 읽기·레벨·품사·한국어 뜻을 표시합니다.'),
    ('문장 입력 (Case B)',
     '문장을 입력하면 형태소별 카드가 나열되고, 레벨 통계와 텍스트 미리보기가 함께 표시됩니다.'),
    ('관리자 모드',
     '패스워드(4649)로 진입 후 읽기·레벨·품사·한국어 뜻을 인라인으로 직접 편집할 수 있습니다. '
     '편집 내용은 localStorage에 자동 저장됩니다.'),
    ('커스텀 단어 등록',
     '사전에 없는 단어(예: 고유명사, 신조어)를 직접 등록하면 이후 분석에서 인식됩니다.'),
    ('사전 CSV 다운로드',
     '관리자가 편집한 단어 목록을 CSV 파일로 내려받아 확인할 수 있습니다.'),
]
for title, desc in screens:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run('■ ')
    r1.font.color.rgb = RGBColor(29,78,216)
    r1.font.size = Pt(11)
    r2 = p.add_run(title)
    set_font(r2, size=11, bold=True, color=(20,20,20))
    add_para(doc, desc, size=10.5, color=(60,60,60), indent=0.5, space_before=0, space_after=6)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 3. 개발 전체 흐름
# ═══════════════════════════════════════════════
add_heading(doc, '3.  개발 전체 흐름', level=1, color=(29,78,216))

add_para(doc, '총 9단계의 프롬프트로 완성. 각 단계마다 Claude Code에 프롬프트를 입력하고 결과를 확인한 뒤 다음 단계로 진행합니다.',
         size=11, space_after=10)

add_screenshot_placeholder(doc, '전체 개발 흐름 다이어그램')

flow_data = [
    ('STEP 1', '프로젝트 구조 생성',        '폴더 생성·데이터 파일 확인',       '5분'),
    ('STEP 2', '한국어 뜻 일괄 번역',       'Claude API로 7,680단어 자동 번역', '20분'),
    ('STEP 3', 'JLPT 사전 빌드',           '중간 JSON 사전 생성',              '3분'),
    ('STEP 4', 'HTML 뼈대·CSS 작성',       'UI 레이아웃·다크테마 완성',         '10분'),
    ('STEP 5', 'Kuromoji 초기화·데이터 내장','형태소 분석기 연결',              '5분'),
    ('STEP 6', '형태소 분석·결과 렌더링',   '핵심 기능 구현',                   '15분'),
    ('STEP 7', '다운로드 기능 구현',        'JSON·CSV·TXT 내보내기',            '5분'),
    ('STEP 8', '통합 테스트·엣지 케이스',   '오류 수정 및 검증',                '10분'),
    ('STEP 9', '빌드 자동화·최종 배포',     '템플릿 분리·README·배포 파일',      '10분'),
]
add_table(doc,
    headers=['단계', '작업 내용', '세부 설명', '소요 시간'],
    rows=flow_data,
    col_widths=[2, 4.5, 7.5, 2.5]
)

add_para(doc, '※ 이후 추가 기능(테마·관리자 모드·커스텀 단어·리팩토링 등)은 위 기본 9단계 완성 후 대화 형식으로 이어서 개발했습니다.',
         size=10, color=(100,100,100), indent=0.3)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 4. 단계별 상세 설명
# ═══════════════════════════════════════════════
add_heading(doc, '4.  단계별 상세 설명', level=1, color=(29,78,216))
add_para(doc, '각 단계에서 Claude Code에 전달한 실제 프롬프트를 포함합니다. '
         '노란 박스가 실제로 입력한 프롬프트 전문입니다.', size=11, color=(80,80,80))

# ──────────────── STEP 1 ────────────────
add_step_box(doc, 1, '프로젝트 구조 생성 및 데이터 확인', '폴더를 만들고 원본 단어 파일의 형식을 검증합니다.')
add_para(doc, '목표: 개발 전 폴더 구조를 잡고, JLPT 단어 데이터 파일(N1~N5) 5개의 형식을 확인합니다.', size=11, color=(40,40,40), indent=0.3)
add_para(doc, '생성되는 구조:', size=10.5, bold=True, indent=0.3)
add_code_block(doc, '''\
tangoya/
├── data/          ← 원본 단어 파일 (N1~N5_words_naver.txt)
├── build/         ← 빌드 스크립트
└── dist/          ← 최종 배포 파일''')
add_para(doc, '데이터 파일 형식 (CSV, 3컬럼):', size=10.5, bold=True, indent=0.3)
add_code_block(doc, '''\
히라가나,한자표기,레벨
あう,会う,N5
たべる,食べる,N5
うつくしい,美しい,N2''')
add_para(doc, '데이터 규모: N5(511개) + N4(877개) + N3(1,308개) + N2(2,262개) + N1(2,722개) = 합계 7,680개', size=10.5, color=(60,60,60), indent=0.3)
add_prompt_box(doc, '''\
아래 지시에 따라 tangoya 프로젝트의 기본 구조를 만들어줘.

【작업 내용】
1. 현재 디렉토리에 다음 폴더 구조를 생성해:
   tangoya/
   ├── data/          (원본 단어 파일을 이동할 폴더)
   ├── build/         (빌드 스크립트를 위한 폴더)
   └── dist/          (최종 배포 파일 tangoya.html이 생성될 폴더)

2. 현재 디렉토리의 N1_words_naver.txt ~ N5_words_naver.txt 파일을
   tangoya/data/ 폴더로 복사해.

3. 각 파일을 읽어 아래 항목을 터미널에 출력해:
   - 파일명
   - 총 라인 수
   - 첫 3줄 (데이터 형식 샘플)
   - 형식이 "히라가나,한자표기,레벨" 3컬럼인지 확인

【기대 결과】
- 5개 파일 모두 확인되고, 형식이 "あう,会う,N5" 패턴임이 확인된다.
- N5: 511줄, N4: 877줄, N3: 1,308줄, N2: 2,262줄, N1: 2,722줄 (합계 7,680줄)

【완료 확인】
위 구조가 정상 생성되었으면 "STEP 1 완료"를 출력해줘.''')

# ──────────────── STEP 2 ────────────────
add_step_box(doc, 2, '한국어 뜻 일괄 번역', 'Claude API를 활용해 7,680개 일본어 단어에 한국어 뜻을 자동으로 번역합니다.')
add_para(doc, '목표: 원본 파일에는 한국어 뜻이 없습니다. Claude API(claude-haiku)를 호출해 100개씩 배치로 번역하고 korean_dict.json을 생성합니다.',
         size=11, color=(40,40,40), indent=0.3)
add_para(doc, '번역 방식:', size=10.5, bold=True, indent=0.3)
bullets2 = [
    '100개씩 배치 처리 → API 호출 횟수: 약 77회',
    '중간 저장: 배치마다 저장하므로 중단 후 재시작 가능',
    '기존 번역은 건너뜀 (재실행 시 중복 API 호출 방지)',
    '결과: korean_dict.json — {"会う": "만나다", "食べる": "먹다", ...}',
]
for b in bullets2:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(b)
    set_font(run, size=10.5, color=(50,50,50))
add_para(doc, '최종 생성: 7,518개 번역 완료 (97.9% 커버율)', size=10.5, color=(29,78,216), indent=0.3)
add_prompt_box(doc, '''\
이 단계는 원본 단어 파일의 7,680개 단어에 한국어 뜻을 붙여
tangoya/data/korean_dict.json 을 생성하는 것이 목표다.

【배경】
원본 파일(N1~N5_words_naver.txt)에는 한국어 뜻이 없다.
Claude API를 사용해 일본어 단어에 한국어 뜻을 일괄 번역하여 별도 JSON으로 저장한다.
이 파일은 이후 STEP 3에서 JLPT 레벨 정보와 합쳐진다.

【작업: tangoya/build/add_korean.py 작성 및 실행】

스크립트 동작:
  1. tangoya/data/ 의 N5→N4→N3→N2→N1 순서로 전체 단어 목록을 수집
  2. 기존 korean_dict.json 이 있으면 로드 → 이미 번역된 단어는 건너뜀
  3. 미번역 단어를 100개씩 배치로 묶어 Claude API 호출
     모델: claude-haiku-4-5-20251001 / max_tokens: 2000
  4. 응답 파싱 실패 시 해당 배치 건너뜀 (중단 없이 계속 진행)
  5. 배치마다 korean_dict.json 에 중간 저장

【korean_dict.json 스키마】
{ "会う": "만나다", "青い": "파랗다", "学生": "학생" }

【완료 확인】
스크립트를 실행해서 tangoya/data/korean_dict.json 이 생성되고
7,000개 이상이 번역 완료되면 "STEP 2 완료"를 출력해줘.''')

# ──────────────── STEP 3 ────────────────
add_step_box(doc, 3, 'JLPT 사전 빌드 (build_dict.py)', '단어·레벨·한국어 뜻을 합쳐 중간 사전 jlpt_dict.json을 생성합니다.')
add_para(doc, '목표: N1~N5 파일 5개 + korean_dict.json을 합쳐 HTML에 내장할 JLPT_DICT를 만듭니다.',
         size=11, color=(40,40,40), indent=0.3)
add_para(doc, '핵심 규칙:', size=10.5, bold=True, indent=0.3)
bullets3 = [
    'N5 우선 원칙: 같은 단어가 여러 레벨에 있으면 가장 낮은 레벨(N5)로 등록',
    '이중 등록: 한자키("会う")와 히라가나키("あう") 모두 등록 → 빠른 검색',
    '스키마: {"r": 읽기, "l": 레벨, "k": 한국어뜻}',
    '결과: 13,680개 항목 (jlpt_dict.json)',
]
for b in bullets3:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(b)
    set_font(run, size=10.5, color=(50,50,50))
add_code_block(doc, '''\
{
  "会う":  {"r": "あう",      "l": "N5", "k": "만나다"},
  "あう":  {"r": "あう",      "l": "N5", "k": "만나다"},
  "学生":  {"r": "がくせい",  "l": "N5", "k": "학생"},
  "美しい":{"r": "うつくしい","l": "N2", "k": "아름답다"}
}''')
add_prompt_box(doc, '''\
tangoya/build/build_dict.py 스크립트를 작성해줘.

【스크립트 역할】
N1~N5 단어 파일 5개 + korean_dict.json 을 합쳐서
tangoya.html에 내장할 최종 JSON 사전(jlpt_dict.json)을 생성한다.

【처리 규칙】
파일 처리 순서: N5 → N4 → N3 → N2 → N1

각 줄 파싱 후 아래 두 가지 키를 모두 등록:
  한자표기 키: "会う" → {"r": "あう", "l": "N5", "k": "만나다"}
  히라가나 키: "あう" → {"r": "あう", "l": "N5", "k": "만나다"}

키 충돌 시: 먼저 등록된 항목 유지 (N5 우선)

【출력】
- 저장 경로: tangoya/build/jlpt_dict.json
- 완료 후 레벨별 항목 수, 한국어 뜻 커버율, 샘플 5개 출력

【완료 확인】
스크립트를 실행해서 tangoya/build/jlpt_dict.json이 생성되고
한국어 뜻이 포함된 샘플이 출력되면 "STEP 3 완료"를 출력해줘.''')

# ──────────────── STEP 4 ────────────────
add_step_box(doc, 4, 'HTML 뼈대 및 CSS 작성', '다크테마 UI, JLPT 레벨 색상 시스템, 반응형 레이아웃을 완성합니다.')
add_para(doc, '목표: JavaScript 로직 없이 HTML 구조와 CSS만으로 전체 레이아웃을 완성합니다.',
         size=11, color=(40,40,40), indent=0.3)
add_para(doc, '레벨별 색상 코드:', size=10.5, bold=True, indent=0.3)
add_table(doc,
    headers=['레벨', '색상', '의미'],
    rows=[
        ['N1', '#ff4d6d (빨강)',    '최고급'],
        ['N2', '#ff8800 (오렌지)',  '고급'],
        ['N3', '#ffd600 (노랑)',    '중급'],
        ['N4', '#00e676 (초록)',    '초중급'],
        ['N5', '#40c4ff (하늘색)', '초급'],
        ['外',  '#6060a0 (보라)',   'JLPT 미등재'],
        ['文法','#444466 (회색)',   '문법 요소'],
    ],
    col_widths=[2.5, 5, 8.5]
)
add_screenshot_placeholder(doc, 'STEP 4 완료 — HTML 레이아웃 골격')
add_prompt_box(doc, '''\
tangoya/dist/tangoya.html 파일을 새로 만들어줘.
이번 단계에서는 JavaScript 로직 없이 HTML 구조와 CSS만 완성한다.

【HTML 구조 요구사항】
<head> 에 포함할 것:
  - charset UTF-8, viewport meta
  - title: "tangoya | 単語屋"
  - Google Fonts: Noto Serif JP (400, 700), Noto Sans KR (300,400,500,700), DM Mono
  - Kuromoji.js CDN: https://cdn.jsdelivr.net/npm/kuromoji@0.1.2/build/kuromoji.js
  - 모든 CSS는 <style> 태그 안에 인라인으로 작성

<body> 구성 (아래 순서로):
  ① 헤더: 배지 "● TANGOYA · 単語屋" + 타이틀 "日本語レベル判定"
  ② 레벨 범례: N1~N5·外 각각 컬러 점 + 텍스트 (가로 나열)
  ③ 입력 카드: textarea(id="inputText", maxlength=1000) + 분석하기 버튼
  ④ 에러 메시지 영역 (id="errorMsg", 초기 hidden)
  ⑤ 로딩 표시 (id="loading", 스피너 + "형태소 분석 중..." 텍스트)
  ⑥ 결과 영역 (id="resultArea") + 다운로드 버튼 3개

【CSS 변수】
  --bg:#0f0f13  --surface:#18181f  --border:#2e2e40  --text:#e8e8f0
  --N1:#ff4d6d  --N2:#ff8800  --N3:#ffd600  --N4:#00e676  --N5:#40c4ff

반응형: 480px 이하에서 버튼 full-width, flex-direction:column

【완료 확인】
브라우저에서 tangoya.html을 열었을 때 레이아웃 골격이 보이면 "STEP 4 완료"를 출력해줘.''')

# ──────────────── STEP 5 ────────────────
add_step_box(doc, 5, 'Kuromoji 초기화 및 사전 데이터 내장', '형태소 분석기를 연결하고 13,680개 사전을 HTML에 내장합니다.')
add_para(doc, 'Kuromoji.js는 JavaScript용 일본어 형태소 분석 라이브러리입니다 (MeCab · IPAdic 기반). '
         'CDN에서 로드되며, 사전 파일을 XHR로 비동기 다운로드합니다.',
         size=11, color=(40,40,40), indent=0.3)
add_code_block(doc, '''\
// Kuromoji 초기화 (페이지 로드 시 자동 실행)
kuromoji.builder({ dicPath: 'https://cdn.jsdelivr.net/npm/kuromoji@0.1.2/dict' })
  .build((err, tokenizer_) => {
    if (err) { initFailed = true; return; }
    tokenizer = tokenizer_;
    console.log('tangoya 준비 완료:', Object.keys(JLPT_DICT).length, '개 단어');
  });''')
add_prompt_box(doc, '''\
tangoya/dist/tangoya.html의 <script> 섹션에 아래를 추가해줘.

【작업 1: JLPT_DICT 데이터 내장】
tangoya/build/jlpt_dict.json 파일을 읽어서
HTML <script> 내부 맨 위에 다음 형식으로 삽입:
  const JLPT_DICT = {전체 JSON 내용};

【작업 2: Kuromoji 초기화】
  - kuromoji.builder({ dicPath: 'https://cdn.jsdelivr.net/npm/kuromoji@0.1.2/dict' })
  - Promise 래핑, 비동기 처리
  - 성공: 전역 tokenizer 에 저장
  - 실패: initFailed = true, 버튼 비활성화, 에러 표시
  - 페이지 로드 시 자동 실행
  - 성공 시: console.log('tangoya 준비 완료:', Object.keys(JLPT_DICT).length, '개 단어')

【작업 3: 보조 함수】
  toKatakana(str): 히라가나 → 카타카나 변환 (+0x60)
  lookupWord(surface, baseForm, reading): JLPT_DICT 검색, 반환값 {r,l,k} 또는 null
    검색 순서: baseForm → surface → reading → ·분리 → 카타카나 변환
  showError(msg) / hideError() / showLoading(bool)

【완료 확인】
브라우저 콘솔에 "tangoya 준비 완료: NNNNN 개 단어" 가 출력되면 "STEP 5 완료"를 출력해줘.''')

# ──────────────── STEP 6 ────────────────
add_step_box(doc, 6, '형태소 분석 및 결과 렌더링', '핵심 기능: 텍스트 → 형태소 분석 → 레벨 판정 → 화면 표시')
add_para(doc, '가장 핵심적인 단계입니다. analyze() 함수가 전체 파이프라인을 처리합니다.',
         size=11, color=(40,40,40), indent=0.3)
add_para(doc, '분석 파이프라인:', size=10.5, bold=True, indent=0.3)
add_code_block(doc, '''\
사용자 입력
  ↓ 언어 검증 (일본어 포함 여부 확인)
  ↓ tokenizer.tokenize(text)  ← Kuromoji 형태소 분석
  ↓ 각 토큰: surface / baseForm / reading / pos 추출
  ↓ lookupWord() → JLPT_DICT 검색 (한자키 → 히라가나키 → 카타카나 순)
  ↓ 레벨 판정: 文法품사 → "文法", 사전히트 → N1~N5, 미등재 → "外"
  ↓ showResult(): 토큰 1개 → Case A (단어 카드) / 2개+ → Case B (문장 카드)''')
add_para(doc, '결과 렌더링 2가지 모드:', size=10.5, bold=True, indent=0.3)
add_table(doc,
    headers=['모드', '조건', '표시 방식'],
    rows=[
        ['Case A', '형태소 1개 (단어 입력)', '대형 카드: 읽기·레벨·품사·사전형·한국어 뜻'],
        ['Case B', '형태소 2개 이상 (문장)', '미니 카드 스트림 + 레벨 통계 + 텍스트 미리보기'],
    ],
    col_widths=[2.5, 5, 8.5]
)
add_screenshot_placeholder(doc, 'STEP 6 완료 — "私は学生です" 분석 결과')
add_prompt_box(doc, '''\
tangoya/dist/tangoya.html의 <script>에
analyze() 함수와 showResult() 함수를 구현해줘.

【analyze() 함수】
트리거: 분석 버튼 클릭, textarea Enter 키 (Shift+Enter 제외)

처리 흐름:
  1. initFailed → 에러, tokenizer 없음 → "사전 로딩 중" 에러
  2. 입력값 trim() → 빈 값이면 에러 "텍스트를 입력해주세요"
  3. hideError(), showLoading(true), 결과 영역 숨김
  4. setTimeout(..., 50) 으로 비동기 처리:
     ① tokenizer.tokenize(text) 실행
     ② 각 토큰에서 surface / baseForm / reading(히라가나변환) / pos / posDetail 추출
     ③ lookupWord(surface, baseForm, readingHira) 호출
     ④ 결과: 사전히트→level=info.l/korean=info.k,
             助詞·助動詞·記号·接続詞→level="文法"/korean="-",
             기타→level="外"/korean="-"
     ⑤ lastResult 저장: {input, tokens, analyzedAt}
     ⑥ showResult(tokens, text) 호출

【showResult() 함수】
tokens.length === 1 → Case A (단어 단독 대형 카드)
tokens.length >= 2  → Case B (문장 미니 카드 스트림 + 레벨 통계)

Case B 구성:
  ① 레벨별 통계 칩 (N5→N1→外→文法, 1개 이상인 레벨만 표시)
  ② 헤더: "분석 결과" + "N 형태소"
  ③ 토큰 카드: 읽기 / 표층형 / 레벨 / 한국어뜻 (없으면 품사)
  ④ 텍스트 미리보기

【필수 포함】
function escHtml(str) {
  return (str||'').replace(/&/g,'&amp;').replace(/</g,'&lt;')
                  .replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

【완료 확인】
"会う" → 한국어 뜻 "만나다" 표시,
"私は学生です" → 각 카드에 한국어 뜻 표시되면 "STEP 6 완료"를 출력해줘.''')

# ──────────────── STEP 7 ────────────────
add_step_box(doc, 7, '다운로드 기능 구현', '분석 결과를 JSON·CSV·TXT 파일로 내보냅니다.')
add_para(doc, '분석 결과를 3가지 형식으로 다운로드할 수 있습니다. 모두 UTF-8 with BOM으로 저장됩니다.',
         size=11, color=(40,40,40), indent=0.3)
add_table(doc,
    headers=['파일', '형식', '포함 내용'],
    rows=[
        ['tangoya_result.json', 'JSON (들여쓰기)', 'input, analyzed_at, tokens배열 (surface/reading/level/korean 등)'],
        ['tangoya_result.csv',  '탭 구분 CSV',     '원문, 읽기, 사전형, 품사, JLPT레벨, 한국어뜻'],
        ['tangoya_result.txt',  '텍스트',          '입력문 + 분석: 단어[읽기, 품사, 레벨, 한국어뜻] 형식'],
    ],
    col_widths=[4.5, 3.5, 8]
)
add_prompt_box(doc, '''\
tangoya/dist/tangoya.html의 <script>에 다운로드 함수 3개를 구현해줘.

【공통】
- lastResult 전역변수를 데이터 소스로 사용 (null이면 아무것도 안 함)
- UTF-8 with BOM:
    new Blob(['\uFEFF' + content], { type: mimeType + ';charset=utf-8;' })
- 공통 헬퍼:
    function downloadFile(content, filename, mimeType) { ... }

【downloadJSON() — tangoya_result.json】
{
  "input": "...",
  "analyzed_at": "ISO8601",
  "tokens": [{ "surface","reading","base_form","pos","level","korean" }]
}

【downloadCSV() — tangoya_result.csv】
헤더: 원문,읽기,사전형,품사,JLPT레벨,한국어뜻
각 값 큰따옴표 래핑, 내부 " → ""

【downloadTXT() — tangoya_result.txt】
1행: 입력: {input}
2행: 분석: {surface}[{reading}, {pos}, {level}, {korean}] ...

【완료 확인】
3가지 다운로드 버튼이 모두 작동하고
JSON에 "korean" 필드, CSV에 "한국어뜻" 컬럼이 포함되면 "STEP 7 완료"를 출력해줘.''')

# ──────────────── STEP 8 ────────────────
add_step_box(doc, 8, '통합 테스트 및 엣지 케이스 검증', '다양한 입력으로 동작을 검증하고 오류를 수정합니다.')
add_para(doc, '테스트 케이스 목록:', size=10.5, bold=True, indent=0.3)
add_table(doc,
    headers=['입력', '예상 결과', '검증 항목'],
    rows=[
        ['会う',           'N5, 만나다',             'Case A, 한국어 뜻'],
        ['食べる',         'N5, 먹다',               'Case A, 동사'],
        ['私は学生です',    'N5/文法/N5/N5',          'Case B, 문법 요소'],
        ['(빈 입력)',      '오류 메시지 표시',        '유효성 검사'],
        ['xyz',           'Case A, 外, 한국어 -',   '비일본어 처리'],
        ['한국어 입력',    '언어 경고 모달 표시',     '언어 감지'],
    ],
    col_widths=[4, 4, 8]
)
add_prompt_box(doc, '''\
tangoya/dist/tangoya.html을 완성된 상태로 열고
아래 테스트를 순서대로 실행해서 결과를 확인해줘.

【테스트 1: 단어 단독 (Case A) — 한국어 뜻 포함 확인】
입력: 会う   → 기대: N5, 만나다
입력: 食べる  → 기대: N5, 먹다
입력: 美しい  → 기대: N2, 한국어 뜻 표시

【테스트 2: 문장 (Case B)】
입력: 私は学生です
기대: 私→N5/"나",  は→文法,  学生→N5/"학생",  です→N5 또는 文法

入力: 東京は大きい都市です
기대: 각 카드에 레벨 + 한국어 뜻 표시

【테스트 3: 엣지 케이스】
(공백) → 에러 "텍스트를 입력해주세요"
xyz    → 外, 한국어 "-"
한국어 → 언어 경고 모달 표시

【테스트 4: 다운로드 확인】
"私は学生です" 분석 후:
  JSON → "korean": "나", "학생" 값 확인
  CSV → "한국어뜻" 컬럼 + 값 포함 확인
  TXT → "[わたし, 名詞, N5, 나]" 형식 확인

【결과 보고 형식】
  ✅ 테스트 1-1 (会う): PASS — N5, 만나다 표시
  ❌ 테스트 X-X: FAIL — 이유

실패 항목은 즉시 수정 후 재테스트.
모든 PASS 시 "STEP 8 완료"를 출력해줘.''')

# ──────────────── STEP 9 ────────────────
add_step_box(doc, 9, '빌드 자동화 및 최종 배포', '빌드 스크립트·README·배포 파일까지 완성합니다.')
add_para(doc, '목표: 언제든 데이터를 업데이트하고 재빌드할 수 있는 자동화 파이프라인을 구성합니다.',
         size=11, color=(40,40,40), indent=0.3)
add_para(doc, '빌드 파이프라인:', size=10.5, bold=True, indent=0.3)
add_code_block(doc, '''\
# 단어 데이터 수정 후 전체 재빌드
python3 build/add_korean.py   # (선택) 한국어 뜻 재번역
python3 build/build_html.py   # jlpt_dict.json → tangoya.html 생성''')
add_para(doc, '템플릿 시스템:', size=10.5, bold=True, indent=0.3)
add_code_block(doc, '''\
tangoya_template.html   (111 KB)  ← // __JLPT_DICT_PLACEHOLDER__ 포함
         ↓  build_html.py 실행
tangoya.html            (854 KB)  ← 13,680개 사전 내장 완성본''')
add_prompt_box(doc, '''\
아래 작업을 순서대로 진행해줘.

【작업 1: 빌드 자동화 스크립트 — tangoya/build/build_html.py】
이 스크립트 하나를 실행하면 STEP 3~5의 과정이 자동화되어
tangoya/dist/tangoya.html 이 완성된 상태로 재생성된다.

처리 흐름:
  1. data/N1~N5 파일 읽기
  2. data/korean_dict.json 로드 (없으면 경고 후 빈 dict 사용)
  3. JLPT_DICT 생성 (r+l+k 포함, N5 우선)
  4. dist/tangoya_template.html 읽기
  5. // __JLPT_DICT_PLACEHOLDER__ → const JLPT_DICT = {json}; 로 교체
  6. dist/tangoya.html 저장
  7. 완료 통계 출력 (항목 수, 파일 크기, 완료 시각)

【작업 2: 템플릿과 최종본 분리】
  tangoya_template.html: JLPT_DICT 데이터 제거 → 플레이스홀더로 표시한 버전
  tangoya.html: build_html.py 실행 결과물 (데이터 내장 완성본)

【작업 3: README.md 작성】
앱 소개, 사용 방법, 기능, 파일 구조, 기술 스택 포함

【최종 체크리스트】
  □ dist/tangoya.html — 브라우저에서 정상 동작
  □ dist/tangoya_template.html — 존재
  □ build/build_html.py — 실행 시 재생성 가능
  □ data/korean_dict.json — 존재, 7,000개 이상
  □ README.md — 존재

모두 확인되면 "STEP 9 완료 — tangoya 빌드 완료"를 출력해줘.''')

doc.add_page_break()

# ═══════════════════════════════════════════════
# 5. 기본 완성 후 추가 개발
# ═══════════════════════════════════════════════
add_heading(doc, '5.  기본 완성 후 추가 개발 (대화형 확장)', level=1, color=(29,78,216))
add_para(doc, '9단계 기본 완성 후, 대화 형식으로 기능을 계속 추가했습니다. '
         '각 요청은 한두 문장의 자연어로 전달했습니다.', size=11)

extras = [
    ('다크/라이트 모드 토글',
     '🌙 버튼으로 테마 전환. CSS 변수를 활용해 전체 색상이 한 번에 바뀝니다. '
     'localStorage에 테마가 저장되어 재방문 시에도 유지됩니다.',
     '다크모드와 라이트모드를 토글하는 버튼 추가해줘.\n'
     '다크모드가 기본이고, 선택한 테마는 localStorage에 저장되게.\n'
     '버튼 위치는 우상단 툴바에 🌙 / ☀️ 이모지로.'),
    ('관리자 모드',
     '🔒 버튼 → 패스워드(4649) 입력 → 🔓 관리자 모드 진입. '
     '읽기·레벨·품사·한국어 뜻을 인라인 편집할 수 있으며, '
     '모든 편집은 localStorage에 자동 저장됩니다.',
     '패스워드로 잠긴 관리자 모드를 추가해줘.\n'
     '관리자만 단어 정보(읽기, 레벨, 품사, 한국어 뜻)를 직접 수정할 수 있게.\n'
     '편집 내용은 localStorage에 저장되어 다음 번 분석 시에도 유지되게.'),
    ('형태소 병합 기능',
     '연속된 토큰을 하나로 합칠 수 있습니다. '
     '예: お + 田 → お田 (커스텀 고유명사). 병합 규칙은 입력 텍스트별로 저장됩니다.',
     '특정 토큰 두 개를 합쳐서 하나의 단어로 처리하는 병합 기능 만들어줘.\n'
     '각 토큰 카드 옆에 + 버튼을 두고, 클릭하면 다음 토큰과 합쳐지게.\n'
     '병합된 단어는 JLPT 사전에서 재검색하고, 규칙은 localStorage에 저장되게.'),
    ('커스텀 단어 등록',
     'JLPT 사전에 없는 단어(고유명사, 신조어 등)를 직접 등록할 수 있습니다. '
     '등록된 단어는 문장 분석 시에도 자동으로 인식됩니다.',
     '사전에 없는 단어를 직접 추가하는 기능을 만들어줘.\n'
     '단어(일본어), 읽기(히라가나), 레벨(N1~N5/外), 한국어 뜻을 입력해서 등록.\n'
     '등록한 단어는 분석 시 자동으로 인식되게. localStorage에 저장.'),
    ('오프라인 동작 지원',
     'Kuromoji 사전·폰트를 로컬에 다운로드해 인터넷 없이도 완전히 동작합니다. '
     'Python 로컬 서버(start_server.py)로 실행합니다.',
     '인터넷 없이도 작동하게 해줘.\n'
     'Kuromoji 사전 파일(dict/*.dat.gz)과 Google Fonts를 로컬에 저장하는 방식으로.\n'
     'Python 로컬 서버 스크립트도 같이 만들어줘. 포트 8000, 브라우저 자동 실행.'),
    ('앱 푸터 추가',
     '개발자 정보·버전·사전 갱신일을 표시하는 푸터를 추가했습니다. '
     '단어를 추가·수정하면 사전 갱신일이 자동으로 업데이트됩니다.',
     '앱 하단에 푸터를 추가해줘.\n'
     '내용: 개발자 Jaehyoring · 버전 v1.0 · 사전 갱신일 표시.\n'
     '사전 갱신일은 build_html.py 실행 시 자동으로 오늘 날짜로 업데이트되게.'),
    ('코드 리팩토링',
     '전역 상수(GRAMMAR_POS, CONFIG, REGEX, COLOR_SCHEMES) 추출, '
     'createStore() 팩토리 패턴으로 localStorage 통일, '
     'processTokens() 분리, autoMergeCustomWords() 독립 함수화.',
     '코드를 리팩토링 해줘.\n'
     '(이 한 문장으로 Claude Code가 스스로 구조를 분석하고\n'
     ' 전역 상수 추출 · 팩토리 패턴 도입 · 함수 분리 등을 제안·실행했습니다.)'),
    ('실행 파일 제작 및 배포',
     'macOS용 .app(PyInstaller), Windows용 .vbs(터미널 창 없이 실행)를 제작해 '
     'release/ 폴더에 배포 파일을 구성했습니다.',
     '실행 파일을 터미널이나 명령프롬프트가 열리는 형식이 아닌\n'
     '바로 실행되게 해줘.\n'
     'macOS는 더블클릭으로 실행되는 .app 파일로, Windows는 .vbs 파일로.'),
    ('GitHub 백업',
     'git 저장소 초기화 후 github.com/Jaehyoring/tangoya에 비공개로 백업했습니다.',
     '깃허브에 푸시해줘.'),
]

for title, desc, prompt in extras:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run('▶ ')
    r1.font.color.rgb = RGBColor(29,78,216)
    r1.font.size = Pt(12)
    r2 = p.add_run(title)
    set_font(r2, size=12, bold=True, color=(20,20,20))
    add_para(doc, desc, size=10.5, color=(60,60,60), indent=0.5, space_before=2, space_after=4)
    add_prompt_box(doc, prompt)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 6. 프롬프트 작성 노하우
# ═══════════════════════════════════════════════
add_heading(doc, '6.  프롬프트 작성 노하우', level=1, color=(29,78,216))
add_para(doc, 'tangoya 개발 과정에서 얻은 효과적인 프롬프트 작성 팁입니다.', size=11)

tips = [
    ('📌 구체적인 결과물을 명시하라',
     '모호한 요청보다 "어떤 파일을, 어떤 형식으로, 어떤 내용으로" 생성하는지 적는 것이 효과적입니다.',
     '❌ "사전 만들어줘"\n✅ "N5~N1 파일과 korean_dict.json을 합쳐서 {"r","l","k"} 스키마의 jlpt_dict.json 생성해줘. N5 우선 원칙 적용."'),
    ('📌 【완료 확인】 기준을 포함하라',
     'AI가 언제 작업을 완료했는지 알 수 있도록 명확한 성공 기준을 제시합니다.',
     '콘솔에 "tangoya 준비 완료: 13680 개 단어"가 출력되면 STEP 5 완료로 확인하겠습니다.'),
    ('📌 단계를 작게 쪼개라',
     '너무 많은 것을 한 번에 요청하면 오류가 많아집니다. 의미 있는 단위로 나눠 진행합니다.',
     'STEP 4(HTML/CSS만)와 STEP 5(JS 초기화)를 분리해서\n각 단계 완성 후 확인하며 진행했습니다.'),
    ('📌 오류 메시지를 그대로 전달하라',
     '오류가 발생하면 콘솔 메시지를 복사해서 그대로 붙여넣으면 AI가 바로 수정합니다.',
     '콘솔 오류:\nUncaught TypeError: Cannot read property \'tokenize\' of null\n→ 이 오류 수정해줘.'),
    ('📌 추가 기능은 짧게 요청해도 된다',
     '기본 기능이 완성된 후에는 자연어 한두 문장으로 기능 추가가 가능합니다.',
     '"코드를 리팩토링 해줘."\n"깃허브에 백업해줘."\n"터미널 없이 바로 실행되는 실행 파일 만들어줘."'),
    ('📌 오류 발생 시 — 디버깅 프롬프트 패턴',
     '오류가 발생했을 때 아래와 같이 상황을 설명하면 AI가 즉시 수정합니다.',
     '【Kuromoji 로딩 실패 시】\n'
     '브라우저 콘솔에서 Kuromoji 관련 에러 발생.\n'
     '콘솔 에러 메시지: [오류 전문 붙여넣기]\n'
     '수정해줘.\n\n'
     '【한국어 뜻이 모두 "-" 로 표시될 때】\n'
     'korean_dict.json이 제대로 생성됐는지 확인하고,\n'
     '문제가 있으면 add_korean.py 를 다시 실행해줘.\n\n'
     '【특정 단어 레벨이 "外" 로 표시될 때】\n'
     '"東京" 이 사전에 있는지 확인해줘. 없으면 커스텀 단어로 등록하는 방법 알려줘.'),
]

for title, desc, ex in tips:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(title)
    set_font(r, size=12, bold=True, color=(29,78,216))
    add_para(doc, desc, size=11, color=(50,50,50), indent=0.3, space_before=2)
    add_code_block(doc, ex)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 7. 기술 스택 및 아키텍처
# ═══════════════════════════════════════════════
add_heading(doc, '7.  기술 스택 및 아키텍처', level=1, color=(29,78,216))

add_heading(doc, '사용 기술', level=2, color=(55,65,81))
add_table(doc,
    headers=['분류', '기술', '역할'],
    rows=[
        ['프론트엔드',   'HTML5 + CSS3 + Vanilla JS', '전체 UI 및 로직 (프레임워크 미사용)'],
        ['형태소 분석',  'Kuromoji.js v0.1.2',        'MeCab·IPAdic 기반 일본어 분석'],
        ['JLPT 데이터',  'HTML 내장 JSON (13,680개)',  '서버 없이 즉시 조회'],
        ['저장',         'localStorage',               '한국어 편집·관리자 편집·테마 저장'],
        ['번역 자동화',  'Claude API (Haiku)',          '7,680개 단어 일괄 번역'],
        ['빌드',         'Python 3',                   'HTML 빌드·오프라인 에셋 다운로드'],
        ['배포 (macOS)', 'PyInstaller (.app)',          '더블클릭 실행 파일'],
        ['배포 (Win)',   'VBScript (.vbs)',             '창 없이 실행'],
        ['버전 관리',    'Git + GitHub',               '소스 코드 백업'],
    ],
    col_widths=[3, 5.5, 7.5]
)

add_heading(doc, '파일 구조 (완성 후)', level=2, color=(55,65,81))
add_code_block(doc, '''\
tangoya/
├── tangoya.app/            ← macOS 실행 파일 (더블클릭)
├── data/
│   ├── N1~N5_words_naver.txt  (7,680개 단어)
│   └── korean_dict.json       (7,518개 한국어 뜻)
├── build/
│   ├── build_html.py       ← HTML 재빌드 자동화
│   ├── build_dict.py       ← 중간 사전 생성
│   ├── add_korean.py       ← 한국어 뜻 자동 번역
│   └── tangoya.spec        ← PyInstaller 빌드 스펙
├── dist/
│   ├── tangoya.html        ← 최종 앱 (854 KB, 사전 내장)
│   ├── tangoya_template.html  ← 빌드용 템플릿
│   ├── start_server.py     ← 로컬 서버 실행기
│   ├── tangoya.vbs         ← Windows 실행 파일
│   ├── kuromoji.js         ← 형태소 분석기
│   ├── dict/               ← Kuromoji 사전 (~18 MB)
│   └── fonts/              ← Google Fonts 로컬 캐시
└── release/                ← 배포용 폴더 (29 MB)
    ├── tangoya.app, tangoya.vbs
    ├── tangoya.html, start_server.py
    ├── kuromoji.js, dict/, fonts/''')

doc.add_page_break()

# ═══════════════════════════════════════════════
# 8. 바이브 코딩 학습 포인트
# ═══════════════════════════════════════════════
add_heading(doc, '8.  바이브 코딩 학습 포인트', level=1, color=(29,78,216))

add_para(doc, '이 프로젝트를 통해 바이브 코딩에서 배울 수 있는 핵심 포인트를 정리합니다.', size=11)

points = [
    ('1️⃣  아이디어 → 실제 동작하는 앱까지',
     '단어 파일 5개와 아이디어만으로 시작해, 약 수 시간 만에 완전히 동작하는 웹 앱을 완성했습니다. '
     '기획(PRD 작성) → 데이터 준비 → UI 구현 → 테스트 → 배포까지 전 과정을 AI와 함께 진행했습니다.'),
    ('2️⃣  단계적 접근의 중요성',
     '9단계로 나눈 접근 방식이 핵심입니다. 각 단계에서 완료 기준을 명확히 하고, '
     '확인 후 다음 단계로 진행함으로써 오류를 조기에 발견하고 수정했습니다.'),
    ('3️⃣  AI는 "코드 조수"가 아닌 "협업 파트너"',
     'AI에게 단순 지시만 하는 것이 아니라, 의도·제약·검증 방법까지 함께 소통합니다. '
     '"왜 이렇게 만들어야 하는지"를 설명하면 AI가 더 좋은 구조를 제안합니다.'),
    ('4️⃣  반복과 개선이 자연스럽다',
     '기본 기능 완성 후 "이것도 추가해줘", "이 부분 개선해줘" 식으로 자연스럽게 기능이 확장됩니다. '
     '일반적인 소프트웨어 개발 사이클(계획→구현→테스트→개선)과 동일한 흐름입니다.'),
    ('5️⃣  완벽하지 않아도 시작할 수 있다',
     '처음부터 모든 기능을 설계하지 않아도 됩니다. '
     '기본 기능을 먼저 만들고, 사용하면서 필요한 기능을 추가하는 방식이 바이브 코딩의 강점입니다.'),
]

for title, desc in points:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(title)
    set_font(r, size=12, bold=True, color=(29,78,216))
    add_para(doc, desc, size=11, color=(50,50,50), indent=0.5, space_before=2, space_after=8)

# ═══════════════════════════════════════════════
# 9. 실습 안내
# ═══════════════════════════════════════════════
add_heading(doc, '9.  직접 해보기 — 실습 안내', level=1, color=(29,78,216))

add_para(doc, '아래 순서로 직접 tangoya 개발 과정을 체험해 볼 수 있습니다.', size=11)

steps_practice = [
    ('준비물 확인',
     ['Python 3.6 이상 설치 (https://www.python.org)',
      'Claude Code 또는 Claude.ai 계정',
      '텍스트 에디터 (VS Code 등)']),
    ('STEP 1~3 실행',
     ['tangoya_claude_code_prompts.md 파일 열기',
      'STEP 1 프롬프트를 Claude Code에 붙여넣기',
      '각 단계 완료 후 다음 단계로 진행']),
    ('앱 실행',
     ['release/tangoya.app (macOS) 또는 release/tangoya.vbs (Windows) 더블클릭',
      '브라우저에서 http://localhost:8000/tangoya.html 확인',
      '"会う" 입력 → 분석 결과 확인']),
    ('기능 추가 도전',
     ['부록 B의 확장 프롬프트 중 하나를 선택해 추가 기능 구현',
      '예: "레벨 필터 버튼 추가" / "분석 이력 표시" / "복사 버튼 추가"']),
]

for i, (step, items) in enumerate(steps_practice, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f'단계 {i}. {step}')
    set_font(r, size=11.5, bold=True, color=(29,78,216))
    for item in items:
        p2 = doc.add_paragraph(style='List Bullet')
        p2.paragraph_format.left_indent = Cm(0.8)
        p2.paragraph_format.space_after = Pt(2)
        run = p2.add_run(item)
        set_font(run, size=10.5, color=(50,50,50))

doc.add_paragraph()
p_note = doc.add_paragraph()
pPr = p_note._element.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'DCFCE7')
pPr.append(shd)
p_note.paragraph_format.left_indent = Cm(0.3)
r = p_note.add_run('✅ 팁: 이미 완성된 tangoya.html이 있으므로 앱을 먼저 사용해 보고, 프롬프트 파일을 참고하여 처음부터 직접 만들어보세요!')
set_font(r, size=11, bold=True, color=(20,100,50))

# 저장
doc.save(OUT)
print(f'✅ 문서 생성 완료: {OUT}')
print(f'   파일 크기: {os.path.getsize(OUT)/1024:.1f} KB')
